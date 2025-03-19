#Sonix.ai .docx editor
#This uses the sonix.ai pre-set docx template and adjusts it according to the MuHSiC conventions
#Created Feb 2025 by Julian Vargo (UC Berkeley)
#Built in Python 3.11.11

#INSTRUCTIONS: Place your docx files into a file folder
#Set your input folder path below
#Set your MuHSiC logo image path below
#Save the document, then run the script

input_folder = r"C:\Users\julia\Downloads\research\muhsic\docx"
image_path = r"C:\Users\julia\Downloads\research\muhsic\docx\Logo_MuHSiC_bicolor.png"

#No need to edit below this line
########
import os
import re
import sys
import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#Make sure LibreOffice is installed before autoconverting to pdf
#Because I couldn't troubleshoot the Mac-compatible section, I used GPT. Email me if it doesn't work.
def export_pdf(file_path):
    abs_path = os.path.abspath(file_path)  # Get absolute path
    output_dir = os.path.dirname(abs_path)  # Save PDF in the same directory

    try:
        if sys.platform.startswith("win"):  # Windows
            soffice_cmd = r"C:\Program Files\LibreOffice\program\soffice.exe"
        else:  # macOS/Linux
            soffice_cmd = "libreoffice"

        # Run LibreOffice to convert .docx to .pdf
        result = subprocess.run(
            [soffice_cmd, "--headless", "--convert-to", "pdf", "--outdir", output_dir, abs_path], 
            capture_output=True, text=True, check=True
        )

        #print(result.stdout)

    except FileNotFoundError:
        print("Error: LibreOffice is not installed or not found. Please install LibreOffice.")
    except subprocess.CalledProcessError as e:
        print(f"LibreOffice conversion failed for {file_path}")
        print(f"Error details: {e.stderr}")

def confirm_proceed():
    while True:
        response = input("Your current documents will be overwritten. Ensure you have a backup. Do you wish to proceed? [y/n]: ").strip().lower()
        if response in ["y", "n"]:
            return response == "y"
        else:
            print("Invalid input. Please rerun and enter 'y' or 'n'.")

#Note that sonix doesn't highlight text. Rather, it "character shades" the text.
def remove_highlight(file_path):
    document = Document(file_path)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            rPr = run._element.find('.//w:rPr', namespaces=document.part.element.nsmap)
            if rPr is not None:
                shading = rPr.find('.//w:shd', namespaces=document.part.element.nsmap)
                if shading is not None:
                    rPr.remove(shading)
    document.save(file_path)

#Recall that this paragraph detects if the speaker code starts with UC, so if another school joins, they'll need to change this script.
def alter_initial_speaker_code(file_path):
    document = Document(file_path)
    first_paragraph = document.paragraphs[0]
    if first_paragraph.text.startswith("UC"):
        for run in first_paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
        first_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    #I'm sure this can be optimized by determining the code based on the file name,
    #I'll come back to this optimization if I find it's a big issue.
    else: 
        print(f"Please add a speaker/file code for {file_path}")
        exit()
    document.save(file_path)

# def correct_first_paragraph_number(file_path):
#     document = Document(file_path)
#     if document.paragraphs:
#         first_paragraph = document.paragraphs[0]
#         text = first_paragraph.text
#         match = re.search(r"_(\d{1,2})_", text)
#         if match:
#             number = match.group(1)
#             if len(number) == 1:
#                 corrected_number = f"_{number.zfill(2)}_"
#                 text = re.sub(r"_(\d{1,2})_", corrected_number, text, count=1)
#                 first_paragraph.text = text
#     document.save(file_path)

def correct_first_paragraph_number(file_path):
    document = Document(file_path)
    filename = os.path.splitext(os.path.basename(file_path))[0]
    if document.paragraphs:
        first_paragraph = document.paragraphs[0]
        if first_paragraph.text.startswith("UC"):
            first_paragraph.clear()
            run = first_paragraph.add_run(filename)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
            first_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    document.save(file_path)

def delete_notes(file_path):
    document = Document(file_path)
    enact_deletion_transcript = False
    enact_deletion_notes = False
    notes_paragraph_number = None
    transcript_paragraph_number = None
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == "Notes":
            notes_paragraph_number = i
            enact_deletion_notes = True
        elif paragraph.text.strip() == "Transcript":
            transcript_paragraph_number = i
            enact_deletion_transcript = True
    if enact_deletion_transcript and enact_deletion_notes:
        for i in range(transcript_paragraph_number - 1, notes_paragraph_number - 1, -1):
            notes_contents_array = document.paragraphs[i]
            notes_contents_array._element.getparent().remove(notes_contents_array._element)
        document.save(file_path)
    else:
        print(f"Warning: 'Notes' or 'Transcript' section missing in {file_path}. Skipping deletion.")

def correct_transcript_paragraph(file_path):
    document = Document(file_path)
    second_paragraph = document.paragraphs[1]
    text = second_paragraph.text
    if "Transcript" in text:
        for run in second_paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = True
        second_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    else:
        new_paragraph = document.add_paragraph()
        new_run = new_paragraph.add_run("Transcript")
        new_run.font.name = "Arial"
        new_run.font.size = Pt(12)
        new_run.bold = True
        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        second_paragraph._element.addprevious(new_paragraph._element)
    document.save(file_path)     

def insert_logo(file_path, image_path):
    document = Document(file_path)
    first_paragraph = document.paragraphs[0] if document.paragraphs else None
    new_paragraph = document.add_paragraph()
    first_paragraph._element.addprevious(new_paragraph._element)
    run = new_paragraph.add_run()
    run.add_picture(image_path, width=Inches(6.5))
    new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.save(file_path)

#For mess-ups when I need to bulk remove the image. Adds some flexibility to the script
def remove_images_from_docx(file_path):
    document = Document(file_path)
    paragraphs_to_remove = []

    # Get namespace mappings from the document
    namespace = {'w': document.part.element.nsmap['w']}  # Fix for namespace issue

    # Scan all paragraphs for images
    for paragraph in document.paragraphs:
        if any(run._element.findall(".//w:drawing", namespace) for run in paragraph.runs):
            paragraphs_to_remove.append(paragraph)

    # Remove identified paragraphs
    for paragraph in paragraphs_to_remove:
        paragraph._element.getparent().remove(paragraph._element)

    # Overwrite the file with modifications
    document.save(file_path)
    print(f"Processed: {os.path.basename(file_path)} (Images removed)")

def process_folder(input_folder):
    for filename in os.listdir(input_folder):
        if filename.startswith("~$") or not filename.endswith(".docx"):
            continue
        file_path = os.path.join(input_folder, filename)
        remove_images_from_docx(file_path)

process_folder(input_folder)

if confirm_proceed():
    print("Beginning docx editing. This may take a while. Do not exit.")
    for filename in os.listdir(input_folder):
        #The Python docx package makes temporary files starting with ~$
        #The following 'if' statement makes sure that these locked temporary files aren't opened in LibreOffice
        if filename.startswith("~$"):
            continue
        if filename.endswith(".docx"):
            file_path = os.path.join(input_folder, filename)
            #delete_notes(file_path)
            remove_highlight(file_path)
            alter_initial_speaker_code(file_path)
            correct_first_paragraph_number(file_path)
            correct_transcript_paragraph(file_path)
            delete_notes(file_path)
            insert_logo(file_path, image_path)
            export_pdf(file_path)
    print("All documents in the folder have been edited")